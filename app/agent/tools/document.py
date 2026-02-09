"""Document text extraction tool (PDF, DOC/DOCX, XLSX)."""
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

MAX_TEXT_LENGTH = 15000  # Max chars to return to avoid LLM token overflow


def _read_pdf(file_path: str) -> str:
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            pages.append(f"[第{i + 1}页]\n{text.strip()}")
    doc.close()

    full = "\n\n".join(pages)
    if len(full) > MAX_TEXT_LENGTH:
        full = full[:MAX_TEXT_LENGTH] + f"\n...[文档共{len(pages)}页，内容已截断]"
    return full


def _read_docx(file_path: str) -> str:
    from docx import Document

    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also read tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(cells))

    full = "\n".join(paragraphs)
    if len(full) > MAX_TEXT_LENGTH:
        full = full[:MAX_TEXT_LENGTH] + "\n...[内容已截断]"
    return full


def _read_xlsx(file_path: str) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(file_path, read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"[工作表: {sheet_name}]")
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                parts.append(" | ".join(cells))
            row_count += 1
            if row_count >= 200:  # Limit rows for large spreadsheets
                parts.append(f"...[仅显示前200行，共更多行]")
                break
    wb.close()

    full = "\n".join(parts)
    if len(full) > MAX_TEXT_LENGTH:
        full = full[:MAX_TEXT_LENGTH] + "\n...[内容已截断]"
    return full


async def read_document(file_path: str) -> str:
    """Read a local document and extract text content.

    Supports PDF, DOC/DOCX, XLSX/XLS formats.
    """
    path = Path(file_path)
    if not path.exists():
        return f"文件不存在: {file_path}"

    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            return _read_pdf(file_path)
        elif ext in (".doc", ".docx"):
            return _read_docx(file_path)
        elif ext in (".xlsx", ".xls"):
            return _read_xlsx(file_path)
        else:
            return f"不支持的文件格式: {ext}"
    except Exception as e:
        logger.error("Failed to read document %s: %s", file_path, e)
        return f"文档读取失败: {e}"
